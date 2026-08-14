# -*- coding: utf-8 -*-
"""rtl_helpers: אכיפת RTL מלאה ב-python-docx, לפי מפרט הסקיל rtl-docx-enforcer.
הקובץ המקורי לא נכלל בהתקנת הסקיל, ולכן נכתב כאן לפי אותה תבנית בדיוק."""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

RTL_STYLES = ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5",
              "Heading 6", "Heading 7", "Heading 8", "Heading 9",
              "List Paragraph", "List Number", "List Bullet", "Title", "Subtitle"]


def _pPr(p):
    pPr = p._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._p.insert(0, pPr)
    return pPr


def _rPr(run):
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._r.insert(0, rPr)
    return rPr


def set_paragraph_rtl(p):
    """bidi ברמת הפסקה. אין jc מפורש: ב-bidi ברירת המחדל start היא ימין."""
    pPr = _pPr(p)
    if pPr.find(qn('w:bidi')) is None:
        pPr.append(OxmlElement('w:bidi'))
    for jc in pPr.findall(qn('w:jc')):
        if jc.get(qn('w:val')) in ('right', 'left'):
            pPr.remove(jc)
    for r in p.runs:
        if not _is_run_explicit_ltr(r):
            set_run_rtl(r)
    return p


def set_run_rtl(run):
    rPr = _rPr(run)
    if rPr.find(qn('w:rtl')) is None:
        el = OxmlElement('w:rtl')
        rPr.append(el)
    return run


def _is_run_explicit_ltr(run):
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        return False
    rtl = rPr.find(qn('w:rtl'))
    return rtl is not None and rtl.get(qn('w:val')) in ('0', 'false')


def set_run_ltr_explicit(run):
    rPr = _rPr(run)
    for el in rPr.findall(qn('w:rtl')):
        rPr.remove(el)
    el = OxmlElement('w:rtl')
    el.set(qn('w:val'), '0')
    rPr.append(el)
    return run


def set_font(run, name="David", size=11, cs_name=None):
    run.font.name = name
    run.font.size = Pt(size)
    rPr = _rPr(run)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), cs_name or name)
    sz = OxmlElement('w:szCs')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    return run


def set_document_rtl(doc):
    """section RTL + כל סגנונות הבסיס מקבלים bidi ברמת הסגנון עצמו."""
    for section in doc.sections:
        sectPr = section._sectPr
        if sectPr.find(qn('w:bidi')) is None:
            sectPr.append(OxmlElement('w:bidi'))
        if sectPr.find(qn('w:rtlGutter')) is None:
            sectPr.append(OxmlElement('w:rtlGutter'))
    for name in RTL_STYLES:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        pPr = st.element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            st.element.append(pPr)
        if pPr.find(qn('w:bidi')) is None:
            pPr.append(OxmlElement('w:bidi'))
        for jc in pPr.findall(qn('w:jc')):
            pPr.remove(jc)
        rPr = st.element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            st.element.append(rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:cs'), 'David')
        if rPr.find(qn('w:rtl')) is None:
            rPr.append(OxmlElement('w:rtl'))
    return doc


def add_rtl_paragraph(doc, text="", style=None, size=11, bold=False, container=None):
    tgt = container if container is not None else doc
    p = tgt.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        set_font(r, size=size)
        set_run_rtl(r)
    set_paragraph_rtl(p)
    return p


def add_rtl_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    set_font(r, size={0: 18, 1: 15, 2: 13, 3: 12}.get(level, 11))
    set_run_rtl(r)
    set_paragraph_rtl(p)
    return p


def add_mixed_rtl_paragraph(doc, parts, size=11, style=None, container=None):
    """parts: [(text, 'rtl'|'ltr'), ...]"""
    tgt = container if container is not None else doc
    p = tgt.add_paragraph(style=style)
    for text, direction in parts:
        r = p.add_run(text)
        set_font(r, size=size)
        if direction == 'ltr':
            set_run_ltr_explicit(r)
        else:
            set_run_rtl(r)
    set_paragraph_rtl(p)
    return p


def add_rtl_numbered_list(doc, items, size=11):
    """מספר ונקודה ב-run נפרד המסומן LTR, כדי שהסדר לא יתהפך."""
    out = []
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{i}. ")
        set_font(r1, size=size)
        set_run_ltr_explicit(r1)
        r2 = p.add_run(item)
        set_font(r2, size=size)
        set_run_rtl(r2)
        set_paragraph_rtl(p)
        pPr = _pPr(p)
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:right'), '360')
        out.append(p)
    return out


def add_rtl_bullet_list(doc, items, size=11):
    out = []
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        set_font(r, size=size)
        set_run_rtl(r)
        set_paragraph_rtl(p)
        pPr = _pPr(p)
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:right'), '360')
        out.append(p)
    return out


def set_table_rtl(table):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    if tblPr.find(qn('w:bidiVisual')) is None:
        tblPr.append(OxmlElement('w:bidiVisual'))
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_paragraph_rtl(p)
                for r in p.runs:
                    if not _is_run_explicit_ltr(r):
                        set_font(r, size=10)
    return table


def final_rtl_audit(doc, verbose=True):
    """רשת ביטחון: עובר על כל המסמך ומתקן במקום לזרוק שגיאה."""
    fixed = {"paragraphs": 0, "tables": 0, "sections": 0, "styles": 0}
    body = doc.element.body
    for p in body.iter(qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)
        if pPr.find(qn('w:bidi')) is None:
            pPr.append(OxmlElement('w:bidi'))
            fixed["paragraphs"] += 1
    for tbl in body.iter(qn('w:tbl')):
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        if tblPr.find(qn('w:bidiVisual')) is None:
            tblPr.append(OxmlElement('w:bidiVisual'))
            fixed["tables"] += 1
    for section in doc.sections:
        sectPr = section._sectPr
        if sectPr.find(qn('w:bidi')) is None:
            sectPr.append(OxmlElement('w:bidi'))
            fixed["sections"] += 1
        if sectPr.find(qn('w:rtlGutter')) is None:
            sectPr.append(OxmlElement('w:rtlGutter'))
    for name in RTL_STYLES[:4]:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        pPr = st.element.find(qn('w:pPr'))
        if pPr is None or pPr.find(qn('w:bidi')) is None:
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                st.element.append(pPr)
            pPr.append(OxmlElement('w:bidi'))
            fixed["styles"] += 1
    if verbose:
        print("RTL audit:", fixed)
    return doc
