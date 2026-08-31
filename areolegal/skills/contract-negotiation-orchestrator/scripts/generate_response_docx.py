"""
generate_response_docx.py

הפקת מסמך התגובה לספק. הסקריפט מקבל את הטיוטה של הספק (עם שינויים שלו),
את ה-JSON של ההכרעות, ומפיק DOCX RTL עם:
    - Track Changes שמשקפים את התגובה של הלקוח (קבלה, דחייה, תיקון).
    - Comments שמסבירים כל הכרעה לפי הפוליסי.
    - עמוד שער עם סיכום ההכרעות.
    - נספח עם טבלת מצב סעיפים (לסבבים מתקדמים).

הסקריפט מסתמך על הסקיל rtl-docx-enforcer.

שימוש:
    python generate_response_docx.py \\
        --supplier-draft draft_from_supplier.docx \\
        --decisions decisions.json \\
        --output response_to_supplier.docx \\
        --round 1
"""

import argparse
import datetime
import json
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET
import shutil
import copy

# יבוא הסקיל
try:
    # איתור rtl_helpers בלי נתיב קשיח: הפלאגין מותקן במקום אחר אצל כל לקוח,
    # ונתיב קבוע כמו /mnt/skills/user/... פשוט אינו קיים שם והסקריפט קורס.
    _here = Path(__file__).resolve()
    _cands = [
        _here.parent.parent.parent / 'rtl-docx-enforcer',   # רכיב אחות באותו פלאגין
        _here.parent.parent / 'rtl-docx-enforcer',
        Path('/mnt/skills/user/rtl-docx-enforcer'),
        Path.home() / '.claude' / 'skills' / 'rtl-docx-enforcer',
    ]
    for _c in _cands:
        if (_c / 'rtl_helpers.py').exists():
            sys.path.insert(0, str(_c))
            break
    from rtl_helpers import (
        set_document_rtl,
        set_paragraph_rtl,
        set_run_rtl,
        set_run_ltr_explicit,
        add_rtl_paragraph,
        add_rtl_heading,
        add_rtl_numbered_list,
        add_rtl_bullet_list,
        set_table_rtl,
        final_rtl_audit,
    )
except ImportError:
    print('שגיאה: הסקיל rtl-docx-enforcer לא טעון.')
    print('הפעל את הסקיל לפני הרצת generate_response_docx.py.')
    sys.exit(1)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# מיפוי קוד החלטה לתיאור
DECISION_LABELS = {
    'A1': 'קבלה מלאה',
    'A2': 'קבלה עם תיקון קל',
    'A3': 'קבלה עם תיקון מהותי',
    'R1': 'דחייה עם הצעה חלופית (Fallback 1)',
    'R2': 'דחייה עם הצעה חלופית (Fallback 2)',
    'R3': 'דחייה מוחלטת - תנאי לחתימה',
}

DECISION_COLORS = {
    'A1': RGBColor(0x00, 0x80, 0x00),  # ירוק כהה
    'A2': RGBColor(0x66, 0x99, 0x00),  # ירוק-זית
    'A3': RGBColor(0xCC, 0x99, 0x00),  # צהוב-זהוב
    'R1': RGBColor(0xCC, 0x66, 0x00),  # כתום
    'R2': RGBColor(0xCC, 0x33, 0x00),  # כתום-אדום
    'R3': RGBColor(0xC0, 0x00, 0x00),  # אדום
}


def add_track_change_insertion(paragraph, text, author='Legal Department', date=None):
    """הוספת טקסט כ-insertion (Track Change).

    יוצר אלמנט w:ins סביב ה-w:r ומסמן עם author ו-date.
    """
    if date is None:
        date = datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')

    # יצירת w:ins
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(_next_track_id()))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date)

    # יצירת w:r עם w:t בתוכו
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rtl = OxmlElement('w:rtl')
    rpr.append(rtl)
    r.append(rpr)

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)

    ins.append(r)
    paragraph._p.append(ins)
    return ins


_track_id_counter = [1000]


def _next_track_id():
    _track_id_counter[0] += 1
    return _track_id_counter[0]


def add_track_change_deletion(paragraph, text, author='Legal Department', date=None):
    """הוספת טקסט כ-deletion (Track Change).

    יוצר אלמנט w:del סביב w:r עם w:delText בתוכו.
    """
    if date is None:
        date = datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')

    d = OxmlElement('w:del')
    d.set(qn('w:id'), str(_next_track_id()))
    d.set(qn('w:author'), author)
    d.set(qn('w:date'), date)

    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rtl = OxmlElement('w:rtl')
    rpr.append(rtl)
    r.append(rpr)

    del_text = OxmlElement('w:delText')
    del_text.text = text
    del_text.set(qn('xml:space'), 'preserve')
    r.append(del_text)

    d.append(r)
    paragraph._p.append(d)
    return d


def add_decision_summary_table(doc, decisions):
    """הוספת טבלה מסכמת של כל ההכרעות בעמוד הראשון."""
    add_rtl_heading(doc, 'סיכום ההכרעות בסבב זה', level=1)

    add_rtl_paragraph(
        doc,
        'הטבלה הבאה מסכמת את ההכרעות לכל הצעת שינוי של הספק. '
        'קוד ההכרעה מצוין בתחילת כל הערת שוליים בגוף המסמך.'
    )
    add_rtl_paragraph(doc, '')

    # ספירת ההכרעות לפי סוג
    counts = {}
    for d in decisions:
        code = d.get('decision_code', 'unknown')
        counts[code] = counts.get(code, 0) + 1

    # טבלת סיכום
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    headers = ['קוד', 'תיאור', 'מספר הופעות']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for code in ['A1', 'A2', 'A3', 'R1', 'R2', 'R3']:
        if code in counts:
            row = table.add_row()
            row.cells[0].text = code
            row.cells[1].text = DECISION_LABELS[code]
            row.cells[2].text = str(counts[code])

    set_table_rtl(table)
    add_rtl_paragraph(doc, '')

    # פירוט ההחלטות
    add_rtl_heading(doc, 'פירוט ההכרעות', level=2)

    detail_table = doc.add_table(rows=1, cols=4)
    detail_table.style = 'Light Grid Accent 1'
    detail_headers = ['#', 'קטגוריה', 'הצעת הספק (תקציר)', 'הכרעה']
    for i, h in enumerate(detail_headers):
        detail_table.rows[0].cells[i].text = h

    for i, dec in enumerate(decisions, 1):
        row = detail_table.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = f'{dec.get("category_number", "?")}. {dec.get("category_name", "")}'
        proposal = dec.get('supplier_proposal_summary', '')
        row.cells[2].text = proposal[:80] + ('...' if len(proposal) > 80 else '')
        code = dec.get('decision_code', '')
        row.cells[3].text = f'[{code}] {DECISION_LABELS.get(code, "")}'

    set_table_rtl(detail_table)


def add_round_status_appendix(doc, decisions, round_number):
    """נספח לסבבים מתקדמים - מצב הסעיפים."""
    if round_number < 2:
        return

    doc.add_page_break()
    add_rtl_heading(doc, f'נספח: מצב הסעיפים בתום סבב {round_number}', level=1)

    # סטטוס לכל סעיף
    add_rtl_paragraph(
        doc,
        'הטבלה מציגה את מצב כל קטגוריה אחרי הסבב הזה. סעיפים בסטטוס STUCK '
        'דורשים שיקול דעת בנוגע להמשך הדרך - האם לעלות לאסקלציה, האם לוותר, או האם '
        'להפסיק את המשא ומתן.'
    )
    add_rtl_paragraph(doc, '')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    headers = ['קטגוריה', 'סטטוס', 'סבבים פתוחים', 'המלצה']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    # קיבוץ ההכרעות לפי קטגוריה
    by_cat = {}
    for d in decisions:
        cat_num = d.get('category_number')
        if cat_num not in by_cat:
            by_cat[cat_num] = []
        by_cat[cat_num].append(d)

    for cat_num, cat_decisions in sorted(by_cat.items()):
        row = table.add_row()
        row.cells[0].text = f'{cat_num}. {cat_decisions[0].get("category_name", "")}'
        # קביעת סטטוס - לפי החלטה הקיצונית ביותר
        codes = [d.get('decision_code', '') for d in cat_decisions]
        if 'R3' in codes:
            status = 'WALK-AWAY'
        elif any(c.startswith('R') for c in codes):
            status = 'STUCK' if round_number > 1 else 'OPEN'
        else:
            status = 'CLOSED'
        row.cells[1].text = status
        row.cells[2].text = str(len([c for c in codes if c.startswith('R')]))

        # המלצה
        if status == 'WALK-AWAY':
            row.cells[3].text = 'דרושה החלטת ניהול - תנאי לחתימה'
        elif status == 'STUCK':
            row.cells[3].text = 'שקול אסקלציה או ויתור'
        else:
            row.cells[3].text = 'המשך לפי המסלול'

    set_table_rtl(table)


def add_decision_comment_paragraph(doc, decision):
    """הוספת פסקת הסבר להכרעה. זוהי גרסה מפושטת של 'הערת שוליים'.

    יצירת comment אמיתי ב-Word דורשת מניפולציה של ה-comments.xml,
    שזה מורכב. הגרסה הזו מוסיפה את ההסבר כפסקת טקסט מודגשת תחת הסעיף.
    """
    code = decision.get('decision_code', '?')
    cat_num = decision.get('category_number', '?')
    cat_name = decision.get('category_name', '')

    # פסקת כותרת ההערה
    header_para = add_rtl_paragraph(doc, '')
    header_run = header_para.add_run(f'[{code}] קטגוריה {cat_num}: {cat_name}')
    set_run_rtl(header_run)
    header_run.font.bold = True
    color = DECISION_COLORS.get(code, RGBColor(0x00, 0x00, 0x00))
    header_run.font.color.rgb = color

    # רכיב א - עמדת הפוליסי
    if decision.get('policy_position'):
        add_rtl_paragraph(doc, f'עמדת המדיניות: {decision["policy_position"]}')

    # רכיב ב - ניתוח הצעת הספק
    if decision.get('supplier_analysis'):
        add_rtl_paragraph(doc, f'ניתוח הצעת הספק: {decision["supplier_analysis"]}')

    # רכיב ג - נימוק משפטי
    if decision.get('legal_reasoning'):
        add_rtl_paragraph(doc, f'נימוק משפטי: {decision["legal_reasoning"]}')

    # רכיב ד - ההכרעה
    if decision.get('resolution'):
        res_para = add_rtl_paragraph(doc, '')
        res_run = res_para.add_run(f'ההכרעה: {decision["resolution"]}')
        set_run_rtl(res_run)
        res_run.font.italic = True

    add_rtl_paragraph(doc, '')


def generate_response_document(supplier_draft_path, decisions_path, output_path,
                               round_number=1, client_name='הלקוח'):
    """הפקת תיק ההכרעות הפנימי.

    supplier_draft_path נשמר בחתימה לשם תאימות אחורה ואינו נקרא: הסקריפט בונה
    מסמך חדש ואינו נוגע בטיוטת הצד שכנגד. סימון ההסכם עצמו בעקוב אחר שינויים
    ובהערות Word אמיתיות נעשה בכלים שב-track_changes_helpers.py.
    """
    if supplier_draft_path:
        print("שים לב: הסקריפט אינו קורא את טיוטת הצד שכנגד ואינו מסמן אותה. "
              "הפלט הוא תיק ההכרעות הפנימי בלבד.", file=sys.stderr)
    # טעינת ההכרעות
    with open(decisions_path, 'r', encoding='utf-8') as f:
        decisions_data = json.load(f)
    decisions = decisions_data.get('decisions', [])

    # יצירת מסמך חדש
    doc = Document()
    set_document_rtl(doc)

    # עמוד שער. הבאנר הראשון ולא הכותרת: המסמך הזה מכיל קודי הכרעה, עמדות
    # מהפלייבוק ונימוקים פנימיים, והוא נקרא בטעות "התייחסות לטיוטת הספק" -- שם
    # שנשמע כמו משהו ששולחים. עורך דין שיעביר אותו ימסור לצד שכנגד את מלוא
    # עמדות המיקוח של החברה, וזה בלתי הפיך.
    warn_para = add_rtl_paragraph(doc, '')
    warn_run = warn_para.add_run('מסמך פנימי — אין להעביר לצד שכנגד')
    set_run_rtl(warn_run)
    warn_run.font.size = Pt(14)
    warn_run.font.bold = True
    warn_run.font.color.rgb = RGBColor(0xB0, 0x00, 0x20)
    sub_para = add_rtl_paragraph(
        doc,
        'המסמך כולל את קודי ההכרעה הפנימיים, את עמדות המדיניות מהפלייבוק, את נוסחי '
        'הנסיגה ואת הנימוקים המשפטיים. הקובץ הנשלח לצד שכנגד הוא ההסכם שלהם בעקוב '
        'אחר שינויים, והוא קובץ אחר.'
    )
    for run in sub_para.runs:
        run.font.italic = True
    add_rtl_paragraph(doc, '')

    title_para = add_rtl_paragraph(doc, '')
    title_run = title_para.add_run(f'סבב {round_number}: תיק ההכרעות הפנימי')
    set_run_rtl(title_run)
    title_run.font.size = Pt(28)
    title_run.font.bold = True

    add_rtl_paragraph(doc, '')
    add_rtl_paragraph(doc, f'לקוח: {client_name}')
    add_rtl_paragraph(doc, f'תאריך: {datetime.date.today().isoformat()}')
    add_rtl_paragraph(doc, f'מספר הכרעות: {len(decisions)}')

    doc.add_page_break()

    # סיכום הכרעות
    add_decision_summary_table(doc, decisions)
    doc.add_page_break()

    # פירוט מלא של ההכרעות
    add_rtl_heading(doc, 'התייחסות מנומקת לכל הצעת שינוי', level=1)
    add_rtl_paragraph(
        doc,
        'בפרק זה מוצגת התייחסות מנומקת לכל הצעת שינוי של הספק. לכל הצעה מצוין הקוד '
        'של ההכרעה, הקטגוריה הרלוונטית במדיניות החוזית, ארבעת רכיבי הנימוק '
        '(עמדת המדיניות, ניתוח הצעת הספק, נימוק משפטי, ההכרעה), והנוסח החלופי '
        'אם רלוונטי.'
    )
    add_rtl_paragraph(doc, '')

    for i, decision in enumerate(decisions, 1):
        add_rtl_heading(doc, f'הצעה #{i}', level=2)

        # ציטוט הצעת הספק
        proposal = decision.get('supplier_proposal_summary', '')
        if proposal:
            quote_para = add_rtl_paragraph(doc, f'הצעת הספק: "{proposal}"')
            for run in quote_para.runs:
                run.font.italic = True

        add_decision_comment_paragraph(doc, decision)

        # נוסח חלופי אם יש
        if decision.get('counter_proposal'):
            add_rtl_paragraph(doc, 'נוסח חלופי מוצע:')
            counter_para = add_rtl_paragraph(doc, decision['counter_proposal'])
            for run in counter_para.runs:
                run.font.bold = True

        add_rtl_paragraph(doc, '')

    # נספח מצב סעיפים לסבבים מתקדמים
    add_round_status_appendix(doc, decisions, round_number)

    # audit סופי
    final_rtl_audit(doc)

    # שמירה
    output_path = Path(output_path)
    doc.save(output_path)
    print(f'נשמר: {output_path}')


def main():
    parser = argparse.ArgumentParser(description='הפקת מסמך תגובה לספק')
    parser.add_argument('--supplier-draft', required=True, help='קובץ DOCX של טיוטת הספק')
    parser.add_argument('--decisions', required=True, help='קובץ JSON של ההכרעות')
    parser.add_argument('--output', required=True, help='קובץ DOCX פלט')
    parser.add_argument('--round', type=int, default=1, help='מספר הסבב')
    parser.add_argument('--client-name', default='הלקוח', help='שם הלקוח')
    args = parser.parse_args()

    generate_response_document(
        args.supplier_draft,
        args.decisions,
        args.output,
        round_number=args.round,
        client_name=args.client_name,
    )


if __name__ == '__main__':
    main()
