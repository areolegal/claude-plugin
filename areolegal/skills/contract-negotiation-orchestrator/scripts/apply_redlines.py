#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""סימון ההסכם של הצד שכנגד בעקוב אחר שינויים אמיתי, עם הערות Word.

זהו התוצר שנשלח לצד שכנגד. הוא נבנה על הקובץ **שלהם**, ולא כמסמך חדש:
השינויים מבוצעים בגוף ההסכם כ-w:ins/w:del, וההסבר נכנס כהערת Word אמיתית.

שני כללי ברזל:

1. **שום תוכן פנימי אינו נכנס לקובץ הזה.** לא קודי הכרעה, לא עמדת הפלייבוק,
   לא נוסח הנסיגה ולא הנימוק המשפטי הפנימי. ההערה לצד שכנגד נלקחת משדה
   ייעודי (`external_comment`) ומשום מקום אחר. תיק ההכרעות הפנימי הוא קובץ נפרד.

2. **לא מנחשים היכן לסמן.** סעיף שלא אותר בוודאות אינו מסומן, והוא מדווח
   כ"לא אותר" כדי שעורך הדין יטפל בו ידנית. סימון שינוי בסעיף הלא נכון
   בהסכם שנשלח לצד שכנגד גרוע בהרבה מסעיף שלא סומן.

שימוש:
    python3 apply_redlines.py --draft <קובץ הצד שכנגד.docx> \\
        --decisions decisions.json --out marked.docx --author "שם עורך הדין"
"""
from __future__ import annotations

import argparse
import difflib
import json
import re
import sys
import unicodedata

try:
    from docx import Document
except ImportError:
    sys.exit("נדרש python-docx. התקנה: pip install python-docx")

import track_changes_helpers as tc

# יחס דמיון מזערי להתאמת ציטוט. נבחר גבוה בכוונה: מוטב לדווח שסעיף לא אותר
# מאשר לסמן שינוי בסעיף שכן. בבדיקות, ניסוח חוזי שונה יורד הרבה מתחת לזה.
MIN_RATIO = 0.82

# מספר סעיף בתחילת פסקה: "5.", "5.1", "5.1.2", "(א)" ועוד
CLAUSE_RE = re.compile(r"^\s*\(?([0-9]+(?:\.[0-9]+)*|[א-ת])\)?[.):]?\s")

RLM = "‏"
BIDI = "".join(chr(c) for c in (0x200E, 0x200F, 0x202A, 0x202B, 0x202C, 0x202D, 0x202E))


def norm(s: str) -> str:
    """נרמול להשוואה: סימני כיווניות, גרשיים, רווחים וניקוד."""
    s = unicodedata.normalize("NFKC", s or "")
    s = s.translate({ord(c): None for c in BIDI})
    s = re.sub(r"[֑-ׇ]", "", s)          # ניקוד וטעמים
    s = s.replace("״", '"').replace("”", '"').replace("“", '"')
    s = s.replace("׳", "'").replace("’", "'").replace("‘", "'")
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def clause_of(text: str) -> str:
    m = CLAUSE_RE.match(text or "")
    return m.group(1) if m else ""


def locate(paragraphs, decision):
    """מאתר את הפסקה שאליה מתייחסת ההכרעה.

    מחזיר (index, method) או (None, reason). הסדר מכוון: ציטוט קודם למספר
    סעיף, מפני שמספור משתנה בין טיוטות ואילו הנוסח נשאר.
    """
    anchor = norm(decision.get("anchor_quote") or decision.get("clause_text") or "")
    want_clause = str(decision.get("clause_number") or "").strip()

    if anchor and len(anchor) >= 12:
        # 1. הכלה מדויקת
        hits = [i for i, p in enumerate(paragraphs) if anchor in norm(p.text)]
        if len(hits) == 1:
            return hits[0], "ציטוט מדויק"
        if len(hits) > 1:
            # ריבוי התאמות: הכרע רק אם מספר הסעיף מפריד ביניהן
            if want_clause:
                narrowed = [i for i in hits if clause_of(paragraphs[i].text) == want_clause]
                if len(narrowed) == 1:
                    return narrowed[0], "ציטוט + מספר סעיף"
            return None, "הציטוט מופיע ב-%d פסקאות ולא ניתן להכריע" % len(hits)

        # 2. התאמה מקורבת, לניסוח שהשתנה קלות
        scored = []
        for i, p in enumerate(paragraphs):
            t = norm(p.text)
            if len(t) < 12:
                continue
            r = difflib.SequenceMatcher(None, anchor, t).ratio()
            if r >= MIN_RATIO:
                scored.append((r, i))
        scored.sort(reverse=True)
        if len(scored) == 1:
            return scored[0][1], "התאמה מקורבת %.2f" % scored[0][0]
        if len(scored) > 1 and scored[0][0] - scored[1][0] >= 0.05:
            return scored[0][1], "התאמה מקורבת %.2f" % scored[0][0]
        if scored:
            return None, "כמה פסקאות דומות במידה זהה; לא ניתן להכריע"

    # 3. נפילה למספר סעיף בלבד
    if want_clause:
        hits = [i for i, p in enumerate(paragraphs) if clause_of(p.text) == want_clause]
        if len(hits) == 1:
            return hits[0], "מספר סעיף"
        if len(hits) > 1:
            return None, "מספר הסעיף %s מופיע %d פעמים" % (want_clause, len(hits))

    if anchor and len(anchor) < 12:
        return None, "הציטוט קצר מדי (%d תווים). נדרשים 12 לפחות" % len(anchor)
    if anchor:
        return None, "הציטוט לא נמצא בטיוטה. ודא שהוא הועתק מנוסח המקור ולא מתמצית"
    return None, ("אין ציטוט. מספר סעיף לבדו אינו מספיק ברוב ההסכמים, "
                  "מפני שהמספור נוצר אוטומטית ב-Word ואינו חלק מהטקסט")


def apply_edit(par, decision, author):
    """מוחק את הנוסח הישן ומכניס את החדש, כשינוי מעקב אמיתי."""
    old = decision.get("original_text") or par.text
    new = decision.get("replacement_text") or ""
    if not new or norm(new) == norm(old):
        return False
    for r in list(par.runs):                 # הנוסח הישן יורד כמחיקה מסומנת
        r._r.getparent().remove(r._r)
    tc.add_del(par, old, ltr=False)
    tc.add_ins(par, new, ltr=False)
    return True


def all_paragraphs(doc):
    """כל הפסקאות במסמך, כולל אלה שבתוך טבלאות.

    `doc.paragraphs` מדלג על תאי טבלה, ובהסכם אמיתי יושבים שם פרטי הצדדים,
    התמורה ולעיתים נספחים שלמים. סעיף שנמצא בטבלה פשוט לא היה נמצא.
    """
    out = list(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                out.extend(cell.paragraphs)
    return out


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--draft", required=True, help="ההסכם של הצד שכנגד")
    ap.add_argument("--decisions", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--author", default="", help="שם עורך הדין, לבועות ההערה")
    ap.add_argument("--report", default="", help="קובץ JSON לדיווח מה אותר ומה לא")
    args = ap.parse_args()

    if args.author:
        tc.AUTHOR = args.author

    data = json.load(open(args.decisions, encoding="utf-8"))
    decisions = data.get("decisions", data if isinstance(data, list) else [])

    doc = Document(args.draft)
    tc.enable_track_changes(doc)
    paragraphs = all_paragraphs(doc)

    applied, commented, missed = [], [], []
    for d in decisions:
        idx, how = locate(paragraphs, d)
        ident = d.get("clause_number") or (d.get("anchor_quote") or "")[:40] or "?"
        if idx is None:
            missed.append({"clause": ident, "reason": how})
            continue
        par = paragraphs[idx]
        changed = apply_edit(par, d, args.author)

        # ההערה לצד שכנגד נלקחת משדה ייעודי בלבד. אין ליפול חזרה על
        # legal_reasoning או policy_position -- אלה פנימיים.
        note = (d.get("external_comment") or "").strip()
        if note:
            tc.add_comment(par, note)
            commented.append(ident)
        if changed:
            applied.append({"clause": ident, "matched_by": how})

    tc.attach_comments(doc)
    doc.save(args.out)

    print("נבנה %s" % args.out)
    print("  שינויים שסומנו : %d" % len(applied))
    print("  הערות לצד שכנגד: %d" % len(commented))
    print("  לא אותרו       : %d" % len(missed))
    for m in missed:
        print("     · %s — %s" % (m["clause"], m["reason"]))
    if missed:
        print("  יש לטפל בהם ידנית. הם לא סומנו, ולא נוחשה עבורם פסקה.")

    if args.report:
        json.dump({"applied": applied, "commented": commented, "missed": missed},
                  open(args.report, "w", encoding="utf-8"), ensure_ascii=False, indent=1)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
