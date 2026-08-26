"""
build_policy_docx.py

בניית מסמך הפוליסי המשפטי כקובץ Word RTL.
הסקריפט מקבל קובץ JSON של פוליסי ומפיק מסמך מובנה ומקצועי.

הסקריפט מסתמך על הסקיל rtl-docx-enforcer לאכיפת RTL.
לפני הפעלה, ודא שהפונקציות מ-rtl_helpers.py זמינות.

שימוש:
    python build_policy_docx.py policy.json --output "מדיניות חוזית - לקוח.docx"
"""

import argparse
import json
import sys
from pathlib import Path

# יבוא הפונקציות RTL מהסקיל - חובה
# אם הסקיל לא טעון, ייכשל בכוונה
try:
    # איתור rtl_helpers בלי נתיב קשיח: הפלאגין מותקן במקום אחר אצל כל לקוח,
    # ונתיב קבוע כמו /mnt/skills/user/... פשוט אינו קיים שם והסקריפט קורס.
    _here = Path(__file__).resolve()
    _cands = [
        _here.parent.parent.parent / 'rtl-docx-enforcer',   # רכיב אחות באותו פלאגין
        _here.parent.parent / 'rtl-docx-enforcer',
        Path('/mnt/skills/user/rtl-docx-enforcer'),
        Path.home() / '.claude' / 'skills' / 'rtl-docx-enforcer',
    ]
    for _c in _cands:
        if (_c / 'rtl_helpers.py').exists():
            sys.path.insert(0, str(_c))
            break
    from rtl_helpers import (
        set_document_rtl,
        set_paragraph_rtl,
        set_run_rtl,
        set_run_ltr_explicit,
        add_rtl_paragraph,
        add_rtl_heading,
        add_rtl_numbered_list,
        add_rtl_bullet_list,
        set_table_rtl,
        final_rtl_audit,
    )
except ImportError:
    print('שגיאה: הסקיל rtl-docx-enforcer לא טעון.')
    print('הפעל את הסקיל לפני הרצת build_policy_docx.py.')
    sys.exit(1)

from docx import Document
from docx.shared import Pt, Cm, RGBColor

from policy_schema import validate_policy, get_category


IMPORTANCE_COLORS = {
    'CRITICAL': RGBColor(0xC0, 0x00, 0x00),  # אדום
    'HIGH': RGBColor(0xE9, 0x6B, 0x00),       # כתום
    'MEDIUM': RGBColor(0x00, 0x66, 0xCC),     # כחול
    'LOW': RGBColor(0x66, 0x66, 0x66),        # אפור
}


def add_cover_page(doc, policy):
    """הוספת עמוד שער."""
    title_para = add_rtl_paragraph(doc, '', style='Title')
    title_run = title_para.add_run('מדיניות חוזית')
    set_run_rtl(title_run)
    title_run.font.size = Pt(36)
    title_run.font.bold = True

    add_rtl_paragraph(doc, '')

    subtitle = add_rtl_paragraph(doc, f'הסכמי שירות - {policy["client_name"]}')
    for run in subtitle.runs:
        run.font.size = Pt(20)
        run.font.bold = True

    add_rtl_paragraph(doc, '')
    add_rtl_paragraph(doc, '')

    info = add_rtl_paragraph(doc, f'גרסה: {policy["policy_version"]}')
    for run in info.runs:
        run.font.size = Pt(14)

    info2 = add_rtl_paragraph(doc, f'תאריך: {policy["policy_date"]}')
    for run in info2.runs:
        run.font.size = Pt(14)

    info3 = add_rtl_paragraph(doc, f'מספר הסכמים שנותחו: {len(policy.get("contracts_analyzed", []))}')
    for run in info3.runs:
        run.font.size = Pt(14)

    # מעבר עמוד
    doc.add_page_break()


def add_introduction(doc, policy):
    """הוספת פרק הקדמה."""
    add_rtl_heading(doc, 'הקדמה', level=1)

    intro_text = policy.get('introduction', '')
    if not intro_text:
        intro_text = (
            'מסמך זה מהווה את המדיניות החוזית של {client_name} ביחס להסכמי שירות '
            'שבהם החברה היא מקבלת השירותים. המסמך נבנה על בסיס ניתוח של {n_contracts} '
            'הסכמים שעליהם חתמה החברה, ומשקף את עמדת המחלקה המשפטית בנושאים החוזיים '
            'המהותיים. המסמך מארגן את העמדה לפי 18 קטגוריות סעיפים סטנדרטיות.'
        ).format(
            client_name=policy['client_name'],
            n_contracts=len(policy.get('contracts_analyzed', [])),
        )

    add_rtl_paragraph(doc, intro_text)
    add_rtl_paragraph(doc, '')

    add_rtl_heading(doc, 'מבנה המסמך', level=2)
    structure_items = [
        'לכל קטגוריה מוצגת עמדת ברירת המחדל - הניסוח המומלץ.',
        'לקטגוריות הקריטיות מוצגים גם Fallback Positions - וויתורים מותרים.',
        'לכל עמדה מצורף נימוק משפטי שמעוגן בדין הישראלי.',
        'בנספחים ניתן למצוא דוגמאות ניסוח מהסכמים קיימים, מטריצת התאמה לסוגי ספקים, '
        'ורשימת ההסכמים שנותחו.',
    ]
    add_rtl_bullet_list(doc, structure_items)

    add_rtl_paragraph(doc, '')

    add_rtl_heading(doc, 'סמכות לאישור חריגות', level=2)
    auth_text = (
        'חריגה מעמדת ברירת המחדל היא בסמכות עורך הדין שמטפל בהסכם. חריגה מ-Fallback 1 '
        'דורשת אישור היועץ המשפטי הראשי או ראש המחלקה המשפטית. חריגה מ-Fallback 2 '
        '(הקו האדום) דורשת אישור בכתב של היועץ המשפטי הראשי וגם של מנכ"ל החברה. '
        'בקטגוריות 6, 7, 9, 11, ו-12 (הקריטיות), חריגה מ-Fallback 2 חייבת תיעוד נפרד.'
    )
    add_rtl_paragraph(doc, auth_text)

    doc.add_page_break()


def add_category_section(doc, category):
    """הוספת פרק לקטגוריה אחת."""
    cat_num = category.get('category_number', 0)
    cat_name = category.get('category_name', '')
    importance = category.get('importance', 'MEDIUM')

    # כותרת הקטגוריה
    heading = add_rtl_heading(doc, f'קטגוריה {cat_num}: {cat_name}', level=1)
    # סימון רמת חשיבות בצבע
    color = IMPORTANCE_COLORS.get(importance, IMPORTANCE_COLORS['MEDIUM'])

    importance_para = add_rtl_paragraph(doc, '')
    imp_run = importance_para.add_run(f'רמת חשיבות: {importance}')
    set_run_rtl(imp_run)
    imp_run.font.color.rgb = color
    imp_run.font.bold = True

    add_rtl_paragraph(doc, '')

    # עמדת ברירת מחדל
    add_rtl_heading(doc, 'עמדת ברירת מחדל', level=2)
    default_pos = category.get('default_position', '')
    if default_pos:
        add_rtl_paragraph(doc, default_pos)
    else:
        warn = add_rtl_paragraph(doc, 'לא הוגדרה עמדת ברירת מחדל. נדרש מילוי.')
        for run in warn.runs:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    add_rtl_paragraph(doc, '')

    # Fallback positions
    fallback_1 = category.get('fallback_1', '')
    if fallback_1:
        add_rtl_heading(doc, 'Fallback 1 - ויתור ראשוני מותר', level=2)
        add_rtl_paragraph(doc, fallback_1)
        add_rtl_paragraph(doc, '')

    fallback_2 = category.get('fallback_2', '')
    if fallback_2:
        add_rtl_heading(doc, 'Fallback 2 - הקו האדום', level=2)
        add_rtl_paragraph(doc, fallback_2)
        add_rtl_paragraph(doc, '')

    # נימוק משפטי
    reasoning = category.get('reasoning', '')
    if reasoning:
        add_rtl_heading(doc, 'נימוק משפטי', level=2)
        add_rtl_paragraph(doc, reasoning)
        add_rtl_paragraph(doc, '')

    # עיגונים בדין
    legal_anchors = category.get('legal_anchors', [])
    if legal_anchors:
        add_rtl_heading(doc, 'עיגונים בדין', level=2)
        add_rtl_bullet_list(doc, legal_anchors)
        add_rtl_paragraph(doc, '')

    # חריגים מתועדים
    exceptions = category.get('documented_exceptions', [])
    if exceptions:
        add_rtl_heading(doc, 'חריגים מתועדים', level=2)
        for exc in exceptions:
            exc_text = f'הסכם: {exc.get("contract", "לא צוין")}. סטייה: {exc.get("deviation", "")}. סיבה: {exc.get("reason", "")}.'
            add_rtl_paragraph(doc, f'• {exc_text}')

    doc.add_page_break()


def add_appendix_supplier_matrix(doc, policy):
    """נספח א - מטריצת ספקים."""
    add_rtl_heading(doc, 'נספח א: מטריצת התאמה לסוגי ספקים', level=1)

    matrix = policy.get('appendix_supplier_matrix', {})
    if not matrix:
        add_rtl_paragraph(doc, 'מטריצה זו אינה זמינה במסמך הנוכחי.')
        return

    # יצירת טבלה: שורה לכל קטגוריה, עמודה לכל סוג ספק
    supplier_types = list(matrix.keys()) if matrix else ['ספק קטן', 'ספק בינוני', 'ספק קריטי', 'ספק בינלאומי']

    table = doc.add_table(rows=1, cols=len(supplier_types) + 1)
    table.style = 'Light Grid Accent 1'

    # כותרות
    header_cells = table.rows[0].cells
    header_cells[0].text = 'קטגוריה'
    for i, sup_type in enumerate(supplier_types):
        header_cells[i + 1].text = sup_type

    # שורות
    for cat in policy.get('categories', []):
        row = table.add_row()
        row.cells[0].text = f'{cat["category_number"]}. {cat["category_name"]}'
        for i, sup_type in enumerate(supplier_types):
            value = matrix.get(sup_type, {}).get(str(cat['category_number']), '')
            row.cells[i + 1].text = value

    # החלת RTL על הטבלה
    set_table_rtl(table)

    doc.add_page_break()


def add_appendix_contracts_list(doc, policy):
    """נספח ב - רשימת הסכמים שנותחו."""
    add_rtl_heading(doc, 'נספח ב: רשימת ההסכמים שנותחו', level=1)

    contracts = policy.get('contracts_analyzed', [])
    if not contracts:
        add_rtl_paragraph(doc, 'אין רשימת הסכמים זמינה.')
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    headers = ['שם הספק', 'תאריך', 'היקף', 'סוג שירות']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for contract in contracts:
        row = table.add_row()
        row.cells[0].text = contract.get('supplier', '')
        row.cells[1].text = contract.get('date', '')
        row.cells[2].text = str(contract.get('value', ''))
        row.cells[3].text = contract.get('service_type', '')

    set_table_rtl(table)


def build_policy_document(policy_path, output_path):
    """בניית המסמך המלא."""
    # טעינה
    policy_path = Path(policy_path)
    with open(policy_path, 'r', encoding='utf-8') as f:
        policy = json.load(f)

    # ולידציה
    errors = validate_policy(policy)
    if errors:
        print('אזהרות בולידציה:')
        for err in errors:
            print(f'  - {err}')
        print()

    # יצירת המסמך
    doc = Document()
    set_document_rtl(doc)

    # בניית התוכן
    add_cover_page(doc, policy)
    add_introduction(doc, policy)

    # פרק לכל קטגוריה
    for category in policy.get('categories', []):
        add_category_section(doc, category)

    # נספחים
    add_appendix_supplier_matrix(doc, policy)
    add_appendix_contracts_list(doc, policy)

    # audit סופי - חובה
    final_rtl_audit(doc)

    # שמירה
    output_path = Path(output_path)
    doc.save(output_path)
    print(f'נשמר: {output_path}')


def main():
    parser = argparse.ArgumentParser(description='בניית מסמך פוליסי משפטי')
    parser.add_argument('policy', help='קובץ JSON של הפוליסי')
    parser.add_argument('--output', required=True, help='קובץ Word פלט')
    args = parser.parse_args()

    build_policy_document(args.policy, args.output)


if __name__ == '__main__':
    main()
